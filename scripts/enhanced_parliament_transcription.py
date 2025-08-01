#!/usr/bin/env python3
"""
Enhanced Automatic Transcription for Parliament TV Videos

This script transcribes Parliament TV videos using speech recognition with
improved memory management and resource handling to prevent timeouts.

Features:
1. Transcribe audio from Parliament TV videos
2. Generate timestamped transcripts
3. Combine with speaker identification for speaker-attributed transcripts
4. Export transcripts in various formats (TXT, SRT, JSON, DOCX)
5. Enhanced memory management and resource cleanup
6. Monitoring of memory usage during transcription
7. Graceful handling of resource constraints

Usage:
    python enhanced_parliament_transcription.py <video_file> [--output OUTPUT_FILE] [--format FORMAT] [--speaker-id SPEAKER_ID]

Example:
    python enhanced_parliament_transcription.py /app/data/temp/parliament_stream_20250501_123045.mp4 --format srt
"""

import os
import sys
import json
import time
import argparse
import logging
import tempfile
import subprocess
import gc
import signal
import threading
import psutil
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Any, Tuple, Union

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger("enhanced_parliament_transcription")

# Constants
WHISPER_MODEL_SIZE = "tiny"  # tiny, base, small, medium, large
LANGUAGE = "en"
SEGMENT_LENGTH = 30  # seconds
OUTPUT_FORMATS = ["txt", "srt", "json", "docx"]
MAX_MEMORY_PERCENT = 80  # Maximum memory usage percentage before forcing cleanup
MEMORY_CHECK_INTERVAL = 5  # Check memory usage every 5 seconds

# Check if we're running in Docker or locally
if os.path.exists("/app"):
    # Docker environment
    TEMP_DIR = Path("/app/data/temp")
    OUTPUT_DIR = Path("/app/data/temp/audio_extracts")
    logger.info("Using Docker paths: %s, %s", TEMP_DIR, OUTPUT_DIR)
else:
    # Local environment
    base_dir = Path(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    TEMP_DIR = base_dir / "data" / "temp"
    OUTPUT_DIR = base_dir / "data" / "temp" / "audio_extracts"
    logger.info("Using local paths: %s, %s", TEMP_DIR, OUTPUT_DIR)

class MemoryMonitor:
    """Class for monitoring memory usage during transcription."""
    
    def __init__(self, max_memory_percent=MAX_MEMORY_PERCENT, check_interval=MEMORY_CHECK_INTERVAL):
        """Initialize the memory monitor."""
        self.max_memory_percent = max_memory_percent
        self.check_interval = check_interval
        self.stop_event = threading.Event()
        self.monitor_thread = None
        self.warning_issued = False
        self.critical_warning_issued = False
    
    def start(self):
        """Start monitoring memory usage."""
        self.stop_event.clear()
        self.warning_issued = False
        self.critical_warning_issued = False
        self.monitor_thread = threading.Thread(target=self._monitor_memory)
        self.monitor_thread.daemon = True
        self.monitor_thread.start()
        logger.info("Memory monitoring started")
    
    def stop(self):
        """Stop monitoring memory usage."""
        if self.monitor_thread:
            self.stop_event.set()
            self.monitor_thread.join(timeout=1.0)
            self.monitor_thread = None
            logger.info("Memory monitoring stopped")
    
    def _monitor_memory(self):
        """Monitor memory usage and trigger cleanup if needed."""
        while not self.stop_event.is_set():
            try:
                # Get current memory usage
                process = psutil.Process(os.getpid())
                memory_info = process.memory_info()
                memory_percent = process.memory_percent()
                memory_mb = memory_info.rss / (1024 * 1024)
                
                # Log memory usage
                logger.info(f"Current memory usage: {memory_mb:.2f} MB ({memory_percent:.2f}%)")
                
                # Check if memory usage is high
                if memory_percent > self.max_memory_percent * 0.9 and not self.warning_issued:
                    logger.warning(f"Memory usage is high ({memory_percent:.2f}%). Consider reducing model size or input length.")
                    self.warning_issued = True
                
                # Check if memory usage is critical
                if memory_percent > self.max_memory_percent and not self.critical_warning_issued:
                    logger.critical(f"Memory usage is critical ({memory_percent:.2f}%). Forcing garbage collection.")
                    self._force_cleanup()
                    self.critical_warning_issued = True
                
                # Sleep for the specified interval
                time.sleep(self.check_interval)
            
            except Exception as e:
                logger.error(f"Error in memory monitor: {e}")
                time.sleep(self.check_interval)
    
    def _force_cleanup(self):
        """Force cleanup of memory."""
        logger.info("Forcing memory cleanup")
        
        # Run garbage collection multiple times
        for _ in range(3):
            gc.collect()
        
        # Try to release memory back to the OS on Linux
        try:
            with open('/proc/sys/vm/drop_caches', 'w') as f:
                f.write('3')
        except (IOError, PermissionError):
            pass

class TranscriptionSegment:
    """Class representing a segment of transcription with timing information."""
    
    def __init__(
        self,
        start: float,
        end: float,
        text: str,
        speaker: Optional[str] = None,
        confidence: float = 1.0
    ):
        self.start = start
        self.end = end
        self.text = text
        self.speaker = speaker
        self.confidence = confidence
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert segment to dictionary."""
        return {
            "start": self.start,
            "end": self.end,
            "text": self.text,
            "speaker": self.speaker,
            "confidence": self.confidence,
            "duration": self.end - self.start
        }
    
    def to_srt(self, index: int) -> str:
        """Convert segment to SRT format."""
        start_time = self._format_time_srt(self.start)
        end_time = self._format_time_srt(self.end)
        
        if self.speaker:
            text = f"[{self.speaker}] {self.text}"
        else:
            text = self.text
        
        return f"{index}\n{start_time} --> {end_time}\n{text}\n"
    
    def to_txt(self) -> str:
        """Convert segment to plain text format."""
        start_time = self._format_time(self.start)
        end_time = self._format_time(self.end)
        
        if self.speaker:
            return f"[{start_time} - {end_time}] {self.speaker}: {self.text}"
        else:
            return f"[{start_time} - {end_time}] {self.text}"
    
    def _format_time(self, seconds: float) -> str:
        """Format time in seconds to HH:MM:SS."""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
    
    def _format_time_srt(self, seconds: float) -> str:
        """Format time in seconds to SRT format (HH:MM:SS,mmm)."""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millisecs = int((seconds - int(seconds)) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{millisecs:03d}"

class EnhancedParliamentTranscriber:
    """Class for transcribing Parliament TV videos with enhanced memory management."""
    
    def __init__(
        self,
        model_size: str = WHISPER_MODEL_SIZE,
        language: str = LANGUAGE,
        segment_length: int = SEGMENT_LENGTH
    ):
        """Initialize the transcriber."""
        self.model_size = model_size
        self.language = language
        self.segment_length = segment_length
        self.temp_dir = TEMP_DIR
        self.output_dir = OUTPUT_DIR
        self.memory_monitor = MemoryMonitor()
        self.model = None
        
        # Create directories if they don't exist
        self.temp_dir.mkdir(parents=True, exist_ok=True)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Set environment variables to limit memory usage
        os.environ["OMP_NUM_THREADS"] = "1"  # Limit OpenMP threads
        os.environ["MKL_NUM_THREADS"] = "1"  # Limit MKL threads
        
        # Check if whisper is installed
        self._check_whisper()
    
    def _check_whisper(self) -> bool:
        """Check if whisper is installed and install if needed."""
        try:
            import whisper
            logger.info(f"Whisper is already installed: {whisper.__version__}")
            return True
        except ImportError:
            logger.info("Whisper not found. Installing...")
            try:
                subprocess.run(
                    [sys.executable, "-m", "pip", "install", "openai-whisper"],
                    check=True
                )
                logger.info("Whisper installed successfully.")
                return True
            except subprocess.CalledProcessError as e:
                logger.error(f"Failed to install whisper: {e}")
                return False
    
    def _load_model(self):
        """Load the Whisper model with memory monitoring."""
        if self.model is not None:
            logger.info("Model already loaded, reusing existing model")
            return
        
        # Force garbage collection before loading model
        self._force_garbage_collection()
        
        # Start memory monitoring
        self.memory_monitor.start()
        
        # Import whisper here to avoid import error if not installed
        import whisper
        
        try:
            logger.info(f"Loading Whisper model: {self.model_size}")
            self.model = whisper.load_model(self.model_size)
            logger.info(f"Whisper model {self.model_size} loaded successfully")
        except Exception as e:
            logger.error(f"Error loading Whisper model: {e}")
            # If we failed to load the requested model, try a smaller one
            if self.model_size not in ["tiny"]:
                fallback_size = "tiny"
                logger.info(f"Attempting to load smaller model: {fallback_size}")
                try:
                    self.model = whisper.load_model(fallback_size)
                    self.model_size = fallback_size
                    logger.info(f"Fallback to {fallback_size} model successful")
                except Exception as e2:
                    logger.error(f"Error loading fallback model: {e2}")
                    raise
    
    def _unload_model(self):
        """Unload the Whisper model and clean up memory."""
        logger.info("Unloading Whisper model and cleaning up memory")
        self.model = None
        self._force_garbage_collection()
    
    def _force_garbage_collection(self):
        """Force garbage collection to free up memory."""
        logger.info("Forcing garbage collection")
        
        # Stop memory monitoring if active
        self.memory_monitor.stop()
        
        # Run garbage collection multiple times
        for _ in range(3):
            gc.collect()
        
        # Log current memory usage
        process = psutil.Process(os.getpid())
        memory_info = process.memory_info()
        memory_percent = process.memory_percent()
        memory_mb = memory_info.rss / (1024 * 1024)
        logger.info(f"Memory usage after cleanup: {memory_mb:.2f} MB ({memory_percent:.2f}%)")
    
    def transcribe(
        self,
        input_path: Path,
        output_path: Optional[Path] = None,
        output_format: str = "txt",
        speaker_id_path: Optional[Path] = None
    ) -> Dict[str, Any]:
        """
        Transcribe an audio or video file with enhanced memory management.
        
        Args:
            input_path: Path to the audio or video file
            output_path: Path to save the transcription (optional)
            output_format: Output format (txt, srt, json, docx)
            speaker_id_path: Path to speaker identification JSON file (optional)
            
        Returns:
            Dict with transcription results
        """
        if not input_path.exists():
            raise FileNotFoundError(f"Input file not found: {input_path}")
        
        # Create output path if not provided
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = self.output_dir / f"transcript_{input_path.stem}_{timestamp}.{output_format}"
        
        # Load speaker identification data if provided
        speaker_data = None
        if speaker_id_path and speaker_id_path.exists():
            try:
                with open(speaker_id_path, 'r') as f:
                    speaker_data = json.load(f)
                logger.info(f"Loaded speaker identification data from {speaker_id_path}")
            except Exception as e:
                logger.error(f"Error loading speaker identification data: {e}")
        
        try:
            # Load the model with memory monitoring
            self._load_model()
            
            # Start memory monitoring for the transcription process
            self.memory_monitor.start()
            
            # Transcribe the audio/video
            logger.info(f"Transcribing file: {input_path}")
            start_time = time.time()
            result = self.model.transcribe(
                str(input_path),
                language=self.language,
                verbose=True
            )
            end_time = time.time()
            logger.info(f"Transcription completed in {end_time - start_time:.2f} seconds")
            
            # Stop memory monitoring
            self.memory_monitor.stop()
            
            # Process the segments
            segments = []
            for i, segment in enumerate(result["segments"]):
                start = segment["start"]
                end = segment["end"]
                text = segment["text"].strip()
                
                # Find speaker if speaker data is available
                speaker = None
                if speaker_data and "timeline" in speaker_data:
                    speaker = self._find_speaker_at_time(speaker_data["timeline"], start)
                
                # Create segment
                segments.append(TranscriptionSegment(
                    start=start,
                    end=end,
                    text=text,
                    speaker=speaker
                ))
            
            # Save the transcription
            self._save_transcription(segments, output_path, output_format)
            
            # Prepare results
            results = {
                "input_file": str(input_path),
                "output_file": str(output_path),
                "model": self.model_size,
                "language": self.language,
                "duration": result["segments"][-1]["end"] if result["segments"] else 0,
                "num_segments": len(segments),
                "has_speaker_data": speaker_data is not None,
                "format": output_format,
                "created_at": datetime.now().isoformat(),
                "processing_time": end_time - start_time
            }
            
            # Save metadata
            metadata_path = output_path.with_suffix('.meta.json')
            with open(metadata_path, 'w') as f:
                json.dump(results, f, indent=2)
            
            logger.info(f"Transcription completed. Output saved to {output_path}")
            
            # Unload model and clean up memory
            self._unload_model()
            
            return results
            
        except Exception as e:
            logger.error(f"Error in transcription: {e}")
            # Ensure memory monitor is stopped and model is unloaded
            self.memory_monitor.stop()
            self._unload_model()
            raise
    
    def _find_speaker_at_time(self, timeline: List[Dict[str, Any]], time: float) -> Optional[str]:
        """Find the speaker at a specific time in the timeline."""
        for entry in timeline:
            if entry["start_time"] <= time <= entry["end_time"]:
                return entry["speaker"]
        return None
    
    def _save_transcription(
        self,
        segments: List[TranscriptionSegment],
        output_path: Path,
        output_format: str
    ) -> None:
        """Save the transcription in the specified format."""
        output_format = output_format.lower()
        
        # Create output directory if it doesn't exist
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        if output_format == "txt":
            with open(output_path, 'w') as f:
                for segment in segments:
                    f.write(segment.to_txt() + "\n")
        
        elif output_format == "srt":
            with open(output_path, 'w') as f:
                for i, segment in enumerate(segments, 1):
                    f.write(segment.to_srt(i) + "\n")
        
        elif output_format == "json":
            with open(output_path, 'w') as f:
                json.dump({
                    "segments": [segment.to_dict() for segment in segments],
                    "metadata": {
                        "created_at": datetime.now().isoformat(),
                        "model": self.model_size,
                        "language": self.language
                    }
                }, f, indent=2)
        
        elif output_format == "docx":
            try:
                from docx import Document
                from docx.shared import Pt
                
                doc = Document()
                doc.add_heading(f"Parliament Transcription", 0)
                
                # Add metadata
                doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
                doc.add_paragraph(f"Model: {self.model_size}")
                doc.add_paragraph(f"Language: {self.language}")
                
                # Add transcription
                doc.add_heading("Transcription", 1)
                for segment in segments:
                    p = doc.add_paragraph()
                    p.add_run(f"[{self._format_time(segment.start)} - {self._format_time(segment.end)}] ").bold = True
                    
                    if segment.speaker:
                        p.add_run(f"{segment.speaker}: ").italic = True
                    
                    p.add_run(segment.text)
                
                doc.save(output_path)
            except ImportError:
                logger.error("python-docx is not installed. Cannot create DOCX file.")
                # Fallback to txt
                output_path = output_path.with_suffix('.txt')
                self._save_transcription(segments, output_path, "txt")
        
        else:
            logger.warning(f"Unsupported format: {output_format}. Using TXT instead.")
            output_path = output_path.with_suffix('.txt')
            self._save_transcription(segments, output_path, "txt")
        
        # Log the output file path in a consistent format that can be parsed by the voice_recognition.py script
        logger.info(f"Transcript saved to: {output_path}")
        print(f"Transcript saved to: {output_path}")
    
    def _format_time(self, seconds: float) -> str:
        """Format time in seconds to HH:MM:SS."""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"

def check_dependencies() -> bool:
    """Check if required dependencies are installed."""
    try:
        # Check for FFmpeg
        subprocess.run(["ffmpeg", "-version"], stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)
        
        # Check for psutil
        try:
            import psutil
            logger.info(f"psutil is installed: {psutil.__version__}")
        except ImportError:
            logger.warning("psutil not installed. Will attempt to install.")
            return False
        
        # Try to import whisper
        try:
            import whisper
            return True
        except ImportError:
            logger.warning("Whisper not installed. Will attempt to install.")
            return False
        
    except subprocess.CalledProcessError:
        logger.error("FFmpeg is not installed. Please install FFmpeg.")
        return False
    except Exception as e:
        logger.error(f"Error checking dependencies: {e}")
        return False

def install_dependencies() -> bool:
    """Install required dependencies."""
    logger.info("Installing dependencies...")
    try:
        # Install whisper and other dependencies
        subprocess.run(
            [sys.executable, "-m", "pip", "install", "openai-whisper", "python-docx", "psutil"],
            check=True
        )
        
        logger.info("Dependencies installed successfully.")
        return True
    except subprocess.CalledProcessError as e:
        logger.error(f"Failed to install dependencies: {e}")
        return False

def setup_signal_handlers():
    """Set up signal handlers for graceful shutdown."""
    def signal_handler(sig, frame):
        logger.info(f"Received signal {sig}, cleaning up and exiting")
        # Force garbage collection
        gc.collect()
        sys.exit(0)
    
    # Register signal handlers
    signal.signal(signal.SIGINT, signal_handler)
    signal.signal(signal.SIGTERM, signal_handler)

def main() -> int:
    """Main entry point for the script."""
    # Set up signal handlers
    setup_signal_handlers()
    
    parser = argparse.ArgumentParser(description='Enhanced Automatic Transcription for Parliament TV Videos')
    parser.add_argument('video_file', help='Path to the audio or video file to transcribe')
    parser.add_argument('--output', '-o', help='Path to save the transcription')
    parser.add_argument('--format', '-f', choices=OUTPUT_FORMATS, default='txt', help='Output format')
    parser.add_argument('--speaker-id', '-s', help='Path to speaker identification JSON file')
    parser.add_argument('--model', '-m', default=WHISPER_MODEL_SIZE, help='Whisper model size')
    parser.add_argument('--language', '-l', default=LANGUAGE, help='Language code')
    parser.add_argument('--input-type', choices=['audio', 'video'], default='video', help='Type of input file')
    parser.add_argument('--max-memory', type=int, default=MAX_MEMORY_PERCENT, 
                        help='Maximum memory usage percentage before cleanup')
    args = parser.parse_args()
    
    try:
        # Check dependencies
        if not check_dependencies():
            if not install_dependencies():
                logger.error("Failed to install dependencies. Exiting.")
                return 1
        
        # Validate input file
        input_path = Path(args.video_file)
        if not input_path.exists():
            logger.error(f"Input file not found: {input_path}")
            return 1
        
        # Create output path if provided
        output_path = None
        if args.output:
            output_path = Path(args.output)
            output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Create speaker ID path if provided
        speaker_id_path = None
        if args.speaker_id:
            speaker_id_path = Path(args.speaker_id)
            if not speaker_id_path.exists():
                logger.warning(f"Speaker ID file not found: {speaker_id_path}")
        
        # Initialize transcriber with enhanced memory management
        transcriber = EnhancedParliamentTranscriber(
            model_size=args.model,
            language=args.language
        )
        
        # Transcribe the file
        result = transcriber.transcribe(
            input_path=input_path,
            output_path=output_path,
            output_format=args.format,
            speaker_id_path=speaker_id_path
        )
        
        logger.info(f"Transcription completed successfully: {result['output_file']}")
        return 0
        
    except Exception as e:
        logger.error(f"Error in transcription: {e}")
        import traceback
        traceback.print_exc()
        return 1

if __name__ == "__main__":
    sys.exit(main())
