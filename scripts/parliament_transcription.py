#!/usr/bin/env python3
"""
Automatic Transcription for Parliament TV Videos

This script transcribes Parliament TV videos using speech recognition and
can optionally combine the transcription with speaker identification data.

Features:
1. Transcribe audio from Parliament TV videos
2. Generate timestamped transcripts
3. Combine with speaker identification for speaker-attributed transcripts
4. Export transcripts in various formats (TXT, SRT, JSON, DOCX)

Usage:
    python parliament_transcription.py <video_file> [--output OUTPUT_FILE] [--format FORMAT] [--speaker-id SPEAKER_ID]

Example:
    python parliament_transcription.py /app/data/temp/parliament_stream_20250501_123045.mp4 --format srt
"""

import os
import sys
import json
import time
import argparse
import logging
import tempfile
import subprocess
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Any, Tuple, Union

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger("parliament_transcription")

# Constants
WHISPER_MODEL_SIZE = "medium"  # tiny, base, small, medium, large
LANGUAGE = "en"
SEGMENT_LENGTH = 30  # seconds
OUTPUT_FORMATS = ["txt", "srt", "json", "docx"]
TEMP_DIR = Path("/app/data/temp")
OUTPUT_DIR = Path("/app/data/media/transcriptions")

class TranscriptionSegment:
    """Class representing a segment of transcription with timing information."""
    
    def __init__(
        self,
        start: float,
        end: float,
        text: str,
        speaker: Optional[str] = None,
        confidence: float = 1.0
    ):
        self.start = start
        self.end = end
        self.text = text
        self.speaker = speaker
        self.confidence = confidence
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert segment to dictionary."""
        return {
            "start": self.start,
            "end": self.end,
            "text": self.text,
            "speaker": self.speaker,
            "confidence": self.confidence,
            "duration": self.end - self.start
        }
    
    def to_srt(self, index: int) -> str:
        """Convert segment to SRT format."""
        start_time = self._format_time_srt(self.start)
        end_time = self._format_time_srt(self.end)
        
        if self.speaker:
            text = f"[{self.speaker}] {self.text}"
        else:
            text = self.text
        
        return f"{index}\n{start_time} --> {end_time}\n{text}\n"
    
    def to_txt(self) -> str:
        """Convert segment to plain text format."""
        start_time = self._format_time(self.start)
        end_time = self._format_time(self.end)
        
        if self.speaker:
            return f"[{start_time} - {end_time}] {self.speaker}: {self.text}"
        else:
            return f"[{start_time} - {end_time}] {self.text}"
    
    def _format_time(self, seconds: float) -> str:
        """Format time in seconds to HH:MM:SS."""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
    
    def _format_time_srt(self, seconds: float) -> str:
        """Format time in seconds to SRT format (HH:MM:SS,mmm)."""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millisecs = int((seconds - int(seconds)) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{millisecs:03d}"

class ParliamentTranscriber:
    """Class for transcribing Parliament TV videos."""
    
    def __init__(
        self,
        model_size: str = WHISPER_MODEL_SIZE,
        language: str = LANGUAGE,
        segment_length: int = SEGMENT_LENGTH
    ):
        """Initialize the transcriber."""
        self.model_size = model_size
        self.language = language
        self.segment_length = segment_length
        self.temp_dir = TEMP_DIR
        self.output_dir = OUTPUT_DIR
        
        # Create directories if they don't exist
        self.temp_dir.mkdir(parents=True, exist_ok=True)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Check if whisper is installed
        self._check_whisper()
    
    def _check_whisper(self) -> bool:
        """Check if whisper is installed and install if needed."""
        try:
            import whisper
            logger.info(f"Whisper is already installed: {whisper.__version__}")
            return True
        except ImportError:
            logger.info("Whisper not found. Installing...")
            try:
                subprocess.run(
                    [sys.executable, "-m", "pip", "install", "openai-whisper"],
                    check=True
                )
                logger.info("Whisper installed successfully.")
                return True
            except subprocess.CalledProcessError as e:
                logger.error(f"Failed to install whisper: {e}")
                return False
    
    def transcribe(
        self,
        video_path: Path,
        output_path: Optional[Path] = None,
        output_format: str = "txt",
        speaker_id_path: Optional[Path] = None
    ) -> Dict[str, Any]:
        """
        Transcribe a video file.
        
        Args:
            video_path: Path to the video file
            output_path: Path to save the transcription (optional)
            output_format: Output format (txt, srt, json, docx)
            speaker_id_path: Path to speaker identification JSON file (optional)
            
        Returns:
            Dict with transcription results
        """
        if not video_path.exists():
            raise FileNotFoundError(f"Video file not found: {video_path}")
        
        # Import whisper here to avoid import error if not installed
        import whisper
        
        # Create output path if not provided
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = self.output_dir / f"transcript_{video_path.stem}_{timestamp}.{output_format}"
        
        # Load speaker identification data if provided
        speaker_data = None
        if speaker_id_path and speaker_id_path.exists():
            try:
                with open(speaker_id_path, 'r') as f:
                    speaker_data = json.load(f)
                logger.info(f"Loaded speaker identification data from {speaker_id_path}")
            except Exception as e:
                logger.error(f"Error loading speaker identification data: {e}")
        
        # Load the model
        logger.info(f"Loading Whisper model: {self.model_size}")
        model = whisper.load_model(self.model_size)
        
        # Transcribe the video
        logger.info(f"Transcribing video: {video_path}")
        result = model.transcribe(
            str(video_path),
            language=self.language,
            verbose=True
        )
        
        # Process the segments
        segments = []
        for i, segment in enumerate(result["segments"]):
            start = segment["start"]
            end = segment["end"]
            text = segment["text"].strip()
            
            # Find speaker if speaker data is available
            speaker = None
            if speaker_data and "timeline" in speaker_data:
                speaker = self._find_speaker_at_time(speaker_data["timeline"], start)
            
            # Create segment
            segments.append(TranscriptionSegment(
                start=start,
                end=end,
                text=text,
                speaker=speaker
            ))
        
        # Save the transcription
        self._save_transcription(segments, output_path, output_format)
        
        # Prepare results
        results = {
            "input_file": str(video_path),
            "output_file": str(output_path),
            "model": self.model_size,
            "language": self.language,
            "duration": result["segments"][-1]["end"] if result["segments"] else 0,
            "num_segments": len(segments),
            "has_speaker_data": speaker_data is not None,
            "format": output_format,
            "created_at": datetime.now().isoformat()
        }
        
        # Save metadata
        metadata_path = output_path.with_suffix('.meta.json')
        with open(metadata_path, 'w') as f:
            json.dump(results, f, indent=2)
        
        logger.info(f"Transcription completed. Output saved to {output_path}")
        return results
    
    def _find_speaker_at_time(self, timeline: List[Dict[str, Any]], time: float) -> Optional[str]:
        """Find the speaker at a specific time in the timeline."""
        for entry in timeline:
            if entry["start_time"] <= time <= entry["end_time"]:
                return entry["speaker"]
        return None
    
    def _save_transcription(
        self,
        segments: List[TranscriptionSegment],
        output_path: Path,
        output_format: str
    ) -> None:
        """Save the transcription in the specified format."""
        output_format = output_format.lower()
        
        if output_format == "txt":
            with open(output_path, 'w') as f:
                for segment in segments:
                    f.write(segment.to_txt() + "\n")
        
        elif output_format == "srt":
            with open(output_path, 'w') as f:
                for i, segment in enumerate(segments, 1):
                    f.write(segment.to_srt(i) + "\n")
        
        elif output_format == "json":
            with open(output_path, 'w') as f:
                json.dump({
                    "segments": [segment.to_dict() for segment in segments],
                    "metadata": {
                        "created_at": datetime.now().isoformat(),
                        "num_segments": len(segments)
                    }
                }, f, indent=2)
        
        elif output_format == "docx":
            try:
                from docx import Document
                from docx.shared import Pt, RGBColor
                
                doc = Document()
                
                # Add title
                title = doc.add_heading("Parliament TV Transcript", level=1)
                
                # Add timestamp
                timestamp = doc.add_paragraph()
                timestamp.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                
                # Add segments
                for segment in segments:
                    p = doc.add_paragraph()
                    
                    # Add timestamp
                    time_run = p.add_run(f"[{self._format_time(segment.start)} - {self._format_time(segment.end)}] ")
                    time_run.font.size = Pt(9)
                    time_run.font.color.rgb = RGBColor(128, 128, 128)
                    
                    # Add speaker if available
                    if segment.speaker:
                        speaker_run = p.add_run(f"{segment.speaker}: ")
                        speaker_run.bold = True
                    
                    # Add text
                    p.add_run(segment.text)
                
                doc.save(output_path)
            except ImportError:
                logger.error("python-docx not installed. Cannot create DOCX file.")
                # Fallback to TXT
                output_path = output_path.with_suffix('.txt')
                self._save_transcription(segments, output_path, "txt")
        
        else:
            logger.warning(f"Unsupported format: {output_format}. Using TXT instead.")
            output_path = output_path.with_suffix('.txt')
            self._save_transcription(segments, output_path, "txt")
    
    def _format_time(self, seconds: float) -> str:
        """Format time in seconds to HH:MM:SS."""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"

def check_dependencies() -> bool:
    """Check if required dependencies are installed."""
    try:
        # Check for FFmpeg
        subprocess.run(["ffmpeg", "-version"], stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)
        
        # Try to import whisper
        try:
            import whisper
            return True
        except ImportError:
            logger.warning("Whisper not installed. Will attempt to install.")
            return False
        
    except subprocess.CalledProcessError:
        logger.error("FFmpeg is not installed. Please install FFmpeg.")
        return False
    except Exception as e:
        logger.error(f"Error checking dependencies: {e}")
        return False

def install_dependencies() -> bool:
    """Install required dependencies."""
    logger.info("Installing dependencies...")
    try:
        # Install whisper
        subprocess.run(
            [sys.executable, "-m", "pip", "install", "openai-whisper", "python-docx"],
            check=True
        )
        
        logger.info("Dependencies installed successfully.")
        return True
    except subprocess.CalledProcessError as e:
        logger.error(f"Failed to install dependencies: {e}")
        return False

def main() -> int:
    """Main entry point for the script."""
    parser = argparse.ArgumentParser(description='Automatic Transcription for Parliament TV Videos')
    parser.add_argument('video_file', help='Path to the video file to transcribe')
    parser.add_argument('--output', '-o', help='Path to save the transcription')
    parser.add_argument('--format', '-f', choices=OUTPUT_FORMATS, default='txt', help='Output format')
    parser.add_argument('--speaker-id', '-s', help='Path to speaker identification JSON file')
    parser.add_argument('--model', '-m', default=WHISPER_MODEL_SIZE, help='Whisper model size')
    parser.add_argument('--language', '-l', default=LANGUAGE, help='Language code')
    args = parser.parse_args()
    
    try:
        # Check dependencies
        if not check_dependencies():
            if not install_dependencies():
                logger.error("Required dependencies could not be installed.")
                return 1
        
        # Create paths
        video_path = Path(args.video_file)
        output_path = Path(args.output) if args.output else None
        speaker_id_path = Path(args.speaker_id) if args.speaker_id else None
        
        # Initialize transcriber
        transcriber = ParliamentTranscriber(
            model_size=args.model,
            language=args.language
        )
        
        # Transcribe the video
        results = transcriber.transcribe(
            video_path=video_path,
            output_path=output_path,
            output_format=args.format,
            speaker_id_path=speaker_id_path
        )
        
        # Print a summary of the results
        print("\nTranscription Results:")
        print(f"Input file: {results['input_file']}")
        print(f"Output file: {results['output_file']}")
        print(f"Duration: {results['duration']:.2f} seconds")
        print(f"Number of segments: {results['num_segments']}")
        print(f"Model: {results['model']}")
        print(f"Format: {results['format']}")
        print(f"Speaker data: {'Yes' if results['has_speaker_data'] else 'No'}")
        
        return 0
    
    except Exception as e:
        logger.error(f"Error: {e}")
        return 1

if __name__ == "__main__":
    sys.exit(main())
